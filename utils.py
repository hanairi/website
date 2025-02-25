from docx import Document

def extract_paper_details(docx_path):
    try:
        doc = Document(docx_path)
        title = doc.paragraphs[0].text if doc.paragraphs else "Untitled Paper"
        author = doc.paragraphs[1].text if len(doc.paragraphs) > 1 else "Unknown Author"
        description = doc.paragraphs[2].text if len(doc.paragraphs) > 2 else "No description available."
        
        return {"title": title, "author": author, "description": description}
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return {"title": "Error", "author": "Error", "description": "Could not read file."}
